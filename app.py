import os
import io
import json
import re
from typing import Literal, Dict, Any, List
import streamlit as st
import pandas as pd
from langchain_core.tools import tool
from langchain_openai import ChatOpenAI
from tavily import TavilyClient
from deepagents import create_deep_agent
import pypdf
import docx

# Step 2: Session state management
if "jobs_df" not in st.session_state:
    st.session_state.jobs_df = None
if "cover_doc" not in st.session_state:
    st.session_state.cover_doc = None
if "last_error" not in st.session_state:
    st.session_state.last_error = ""
if "raw_final" not in st.session_state:
    st.session_state.raw_final = ""

# Step 3: User interface setup
st.set_page_config(page_title="Job Application Assistant", page_icon="💼", layout="wide")
st.title("💼 Job Application Assistant")
c0, c1, c2 = st.columns([2, 1, 1])
with c0:
    uploaded = st.file_uploader("Upload your resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
with c1:
    target_title = st.text_input("Target title", "Senior Machine Learning Engineer")
with c2:
    target_location = st.text_input("Target location(s)", "Bangalore OR Remote")
skills_hint = st.text_area(
    "Add/override skills (optional)",
    "",
    placeholder="Python, PyTorch, LLMs, RAG, Azure, vLLM, FastAPI",
)

# Step 4: File processing helper functions
def extract_text(file) -> str:
    if not file:
        return ""
    name = file.name.lower()
    if name.endswith(".txt"):
        return file.read().decode("utf-8", errors="ignore")
    if name.endswith(".pdf"):
        pdf = pypdf.PdfReader(io.BytesIO(file.read()))
        return "\n".join((p.extract_text() or "") for p in pdf.pages)
    if name.endswith(".docx"):
        d = docx.Document(io.BytesIO(file.read()))
        return "\n".join(p.text for p in d.paragraphs)
    return ""

def md_to_docx(md_text: str) -> bytes:
    doc = docx.Document()
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            doc.add_heading(line.lstrip("#").strip(), level=level)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# Step 5: Data processing and extraction
def normalize_jobs(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    normed = []
    for it in items:
        if not isinstance(it, dict):
            continue
        # case-insensitive keys
        lower_map = {str(k).strip().lower(): it[k] for k in it.keys()}
        company = str(lower_map.get("company", "") or "").strip()
        title = str(lower_map.get("title", "") or "").strip()
        location = str(lower_map.get("location", "") or "").strip()
        link = str(lower_map.get("link", "") or "").strip()
        why_fit = str(lower_map.get("why_fit", lower_map.get("good match", "")) or "").strip()
        if not link:
            continue
        normed.append({
            "company": company or "—",
            "title": title or "—",
            "location": location or "—",
            "link": link,
            "Good Match": why_fit or "—",
        })
    return normed[:5]

def extract_jobs_from_text(text: str) -> List[Dict[str, Any]]:
    if not text:
        return []
    pattern = r"<JOBS>\s*(?:```[\w-]*\s*)?(\[.*?\])\s*(?:```)?\s*</JOBS>"
    m = re.search(pattern, text, flags=re.S | re.I)
    if not m:
        return []
    raw = m.group(1).strip().strip("`").strip()
    try:
        obj = json.loads(raw)
        return obj if isinstance(obj, list) else []
    except Exception:
        try:
            salvaged = re.sub(r"(?<!\\)'", '"', raw)
            obj = json.loads(salvaged)
            return obj if isinstance(obj, list) else []
        except Exception:
            st.session_state.last_error = f"JSON parse failed: {raw[:1200]}"
            return []

# API Key Configuration via Sidebar
st.sidebar.title("🔑 API Configuration")
st.sidebar.markdown("Configure your credentials to run the assistant.")

# Load existing environment variables
env_openai_key = os.environ.get("OPENAI_API_KEY", "")
env_tavily_key = os.environ.get("TAVILY_API_KEY", "")

openai_input = st.sidebar.text_input(
    "OpenAI API Key",
    value=env_openai_key,
    type="password",
    placeholder="sk-proj...",
    help="Required for language understanding and cover letter generation."
)
tavily_input = st.sidebar.text_input(
    "Tavily API Key",
    value=env_tavily_key,
    type="password",
    placeholder="tvly-dev...",
    help="Required for web search functionality to find up-to-date job postings."
)

if openai_input:
    os.environ["OPENAI_API_KEY"] = openai_input
if tavily_input:
    os.environ["TAVILY_API_KEY"] = tavily_input

# Step 6: Tool integration
@tool
def internet_search(
    query: str,
    max_results: int = 5,
    topic: Literal["general", "news", "finance"] = "general",
    include_raw_content: bool = False,
) -> List[Dict[str, Any]]:
    """Search the internet for current job postings and web content using Tavily."""
    tavily_key = os.environ.get("TAVILY_API_KEY", "")
    if not tavily_key:
        raise RuntimeError("TAVILY_API_KEY is not set.")
    client = TavilyClient(api_key=tavily_key)
    return client.search(
        query=query,
        max_results=max_results,
        include_raw_content=include_raw_content,
        topic=topic,
    )

# Step 7: Deep agent configuration
INSTRUCTIONS = (
    "You are a job application assistant. Do two things:\n"
    "1) Use the web search tool to find exactly 5 CURRENT job postings (matching the user's target title, locations, and skills). "
    "Return them ONLY as JSON in this exact wrapper:\n"
    "<JOBS>\n"
    "[{\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"}, ... five total]\n"
    "</JOBS>\n"
    "Rules: The list must be valid JSON (no comments), real links to the job page or application page, no duplicates.\n"
    "2) Produce a concise cover letter (≤150 words) for EACH job, with a subject line, appended to cover_letters.md under a heading per job.\n"
    "Do not invent jobs. Prefer reputable sources (company career pages, LinkedIn, Lever, Greenhouse)."
)
JOB_SEARCH_PROMPT = (
    "Search and select 5 real postings that match the user's title, locations, and skills. "
    "Output ONLY this block format (no extra text before/after the wrapper):\n"
    "<JOBS>\n"
    "[{\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"},"
    " {\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"},"
    " {\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"},"
    " {\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"},"
    " {\"company\":\"...\",\"title\":\"...\",\"location\":\"...\",\"link\":\"https://...\",\"Good Match\":\"one sentence\"}]"
    "\n</JOBS>"
)
COVER_LETTER_PROMPT = (
    "For each job in the found list, write a subject line and a concise cover letter (≤150 words) that ties the user's skills/resume to the role. "
    "Append to cover_letters.md under a heading per job. Keep writing tight and specific."
)

def build_agent():
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        st.error("Please set OPENAI_API_KEY in your environment.")
        st.stop()
    llm = ChatOpenAI(model=os.environ.get("OPENAI_MODEL", "gpt-4o-mini"), temperature=0.2)
    tools = [internet_search]
    subagents = [
        {"name": "job-search-agent", "description": "Finds relevant jobs", "system_prompt": JOB_SEARCH_PROMPT},
        {"name": "cover-letter-writer-agent", "description": "Writes cover letters", "system_prompt": COVER_LETTER_PROMPT},
    ]
    return create_deep_agent(model=llm, tools=tools, system_prompt=INSTRUCTIONS, subagents=subagents)

def make_task_prompt(resume_text: str, skills_hint: str, title: str, location: str) -> str:
    skills = skills_hint.strip()
    skill_line = f" Prioritize these skills: {skills}." if skills else ""
    return (
        f"Target title: {title}\n"
        f"Target location(s): {location}\n"
        f"{skill_line}\n\n"
        f"RESUME RAW TEXT:\n{resume_text[:8000]}"
    )

# Step 8: Main application logic
resume_text = extract_text(uploaded) if uploaded else ""
run_clicked = st.button("Run", type="primary", disabled=not uploaded)
if run_clicked:
    st.session_state.last_error = ""
    st.session_state.raw_final = ""
    try:
        if not os.environ.get("OPENAI_API_KEY"):
            st.error("OPENAI_API_KEY is not set. Please add it to the sidebar configuration or environment variables.")
            st.stop()
        if not os.environ.get("TAVILY_API_KEY"):
            st.error("TAVILY_API_KEY is not set. Please add it to the sidebar configuration or environment variables.")
            st.stop()
        agent = build_agent()
        task = make_task_prompt(resume_text, skills_hint, target_title, target_location)
        state = {
            "messages": [{"role": "user", "content": task}],
        }
        with st.spinner("Finding jobs and drafting cover letters..."):
            result = agent.invoke(state)
        final_msgs = result.get("messages", [])
        final_text = (final_msgs[-1].content if final_msgs else "") or ""
        st.session_state.raw_final = final_text

        # Extract cover letters - try multiple sources:
        cover_md = ""

        # Source 1: Check if the agent wrote cover_letters.md to disk
        cover_file_path = os.path.join(os.getcwd(), "cover_letters.md")
        if os.path.exists(cover_file_path):
            with open(cover_file_path, "r", encoding="utf-8") as f:
                cover_md = f.read().strip()
            # Clean up the file after reading
            os.remove(cover_file_path)

        # Source 2: Extract from write_file tool calls in messages
        if not cover_md:
            for msg in final_msgs:
                if hasattr(msg, "tool_calls"):
                    for tc in msg.tool_calls:
                        if tc.get("name") in ("write_file", "edit_file") and "cover_letter" in str(tc.get("args", {})).lower():
                            args = tc.get("args", {})
                            content = args.get("content", "") or args.get("file_text", "")
                            if content:
                                cover_md += content + "\n\n"

        # Source 3: Extract from the final message text (fallback)
        if not cover_md:
            # Look for cover letter sections in the conversation
            for msg in reversed(final_msgs):
                content = getattr(msg, "content", "") or ""
                if isinstance(content, str) and ("cover letter" in content.lower() or "subject:" in content.lower()):
                    # Check if this looks like actual cover letter content (not just a mention)
                    if len(content) > 200:
                        cover_md = content
                        break

        st.session_state.cover_doc = md_to_docx(cover_md) if cover_md else None
        raw_jobs = extract_jobs_from_text(final_text)
        jobs_list = normalize_jobs(raw_jobs)
        st.session_state.jobs_df = pd.DataFrame(jobs_list) if jobs_list else None
        st.success("Done. Results generated and saved.")
    except Exception as e:
        st.session_state.last_error = str(e)
        st.error(f"Error: {e}")

# Step 9: Results display and download
st.header("Jobs")
if st.session_state.jobs_df is None or st.session_state.jobs_df.empty:
    st.write("No jobs to show yet.")
else:
    df = st.session_state.jobs_df.copy()
    def as_link(u: str) -> str:
        u = u if isinstance(u, str) else ""
        return f'<a href="{u}" target="_blank">Apply</a>' if u else "—"
    if "link" in df.columns:
        df["link"] = df["link"].apply(as_link)
    cols = [c for c in ["company", "title", "location", "link", "Good Match"] if c in df.columns]
    df = df[cols]
    st.write(df.to_html(escape=False, index=False), unsafe_allow_html=True)

st.header("Download")
if st.session_state.cover_doc:
    st.download_button(
        "Download cover_letters.docx",
        data=st.session_state.cover_doc,
        file_name="cover_letters.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="dl_cover_letters",
    )
else:
    st.caption("Cover letters not produced yet.")
